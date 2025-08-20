import io
import os
import json
import google.generativeai as genai
from flask import Flask, request, jsonify, send_file , render_template
from docx import Document
from dotenv import load_dotenv

# Import the Gemini service
from services.gemini_service import get_gemini_analysis

# Load environment variables from .env file
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Get the absolute path of the current script's directory
basedir = os.path.abspath(os.path.dirname(__file__))

# Configure Flask with the correct paths for templates and static files
# templates_folder is where Flask will look for HTML files
# static_folder is where Flask will look for JS, CSS, images, etc.
app = Flask(__name__,
            template_folder=os.path.join(basedir, '../frontend'),
            static_folder=os.path.join(basedir, '../frontend'))
# --- MODIFIED CODE END ---

def extract_text_from_docx(docx_file):
    """
    Extracts text from a .docx file.
    """
    doc = Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def create_resume_docx(resume_content):
    """
    Creates a new .docx file with a standardized style from a string.
    """
    document = Document()
    document.add_heading(resume_content.split('\n')[0], level=0)
    lines = resume_content.split('\n')
    for line in lines[1:]:
        if line.strip().isupper() and len(line) < 50:
            document.add_heading(line.strip(), level=1)
        elif line.strip().startswith('-'):
            document.add_paragraph(line.strip(), style='List Bullet')
        else:
            document.add_paragraph(line.strip())
            
    temp_file = io.BytesIO()
    document.save(temp_file)
    temp_file.seek(0)
    return temp_file

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    if 'resume' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['resume']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file and file.filename.endswith('.docx'):
        try:
            resume_text = extract_text_from_docx(file)
            analysis_result = get_gemini_analysis(resume_text)
            
            if "error" in analysis_result:
                return jsonify(analysis_result), 500
            
            return jsonify(analysis_result), 200
            
        except Exception as e:
            return jsonify({"error": f"An error occurred during file processing: {e}"}), 500
            
    return jsonify({"error": "Invalid file type. Please upload a .docx file."}), 400

@app.route('/download_updated_resume', methods=['POST'])
def download_updated_resume():
    data = request.json
    updated_resume_text = data.get('updated_resume_text')
    
    if not updated_resume_text:
        return jsonify({"error": "Updated resume content not found."}), 400
    
    try:
        updated_resume_file = create_resume_docx(updated_resume_text)
        return send_file(
            updated_resume_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='updated_resume.docx'
        )
    except Exception as e:
        return jsonify({"error": f"An error occurred while creating the file: {e}"}), 500

if __name__ == '__main__':
    app.run(debug=True)